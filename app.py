import streamlit as stimport os
token = os.getenv("HF_API_KEY")

import requests
import plotly.express as px
import os
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
import json
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
import io
from PIL import Image
from dotenv import load_dotenv

# --- Enhanced CSS Styling ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@600;700&display=wrap');

    .stApp {
        background-image: url('https://images.unsplash.com/photo-1509316975850-ff9c5deb0cd9?ixlib=rb-4.0.3&auto=format&fit=crop&w=2070&q=80');
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        background-attachment: fixed;
        filter: brightness(0.85);
        min-height: 100vh;
    }

    main > div.block-container {
        background-color: rgba(255, 255, 255, 0.9);
        padding: 20px 30px 30px 30px;
        border-radius: 12px;
    }

    .header-footer {
        text-align: center;
        padding: 20px 10px;
        background: linear-gradient(to right, #D2691E, #F4A460);
        color: white;
        border-radius: 12px;
        margin-bottom: 20px;
        font-family: 'Poppins', sans-serif;
        font-weight: 700;
    }

    .footer {
        margin-top: 40px;
        font-size: 0.9em;
        margin-bottom: 0;
    }

    .typewriter {
        font-size: 40px;
        font-family: 'Poppins', sans-serif;
        font-weight: 700;
        overflow: hidden;
        white-space: nowrap;
        animation: typing 3.5s steps(40, end) forwards;
        border-right: none !important;
        margin-bottom: 10px;
    }

    @keyframes typing {
        from { width: 0; }
        to { width: 100%; }
    }

    /* Blue button style used for Analyze and Export Report */
    .stButton > button {
        background-color: #007BFF !important;
        color: white !important;
        border: none !important;
        padding: 10px 20px !important;
        font-weight: bold !important;
        border-radius: 8px !important;
        width: 100% !important;
        transition: 0.3s ease !important;
        cursor: pointer !important;
        font-family: 'Poppins', sans-serif;
        font-size: 18px;
    }

    .stButton > button:hover {
        background-color: #0056b3 !important;
    }
    /* Style for the new "Clear All" button if it uses a different class */
    .stButton > button.clear-button {
        background-color: #6c757d !important; /* Muted gray */
    }
    .stButton > button.clear-button:hover {
        background-color: #5a6268 !important;
    }

    /* Sentiment label box */
    .sentiment-box {
        border: 2px solid #007BFF;
        background-color: #e7f0fe;
        padding: 15px 20px;
        border-radius: 12px;
        font-family: 'Poppins', sans-serif;
        font-weight: 700;
        font-size: 24px;
        color: #007BFF;
        margin: 15px 0;
        text-align: center;
    }

    /* AI comment blockquote */
    .ai-comment {
        border-left: 4px solid #764ba2;
        background-color: #f0e9ff;
        padding: 10px 15px;
        margin: 20px 0 20px 0;
        font-size: 18px;
        font-family: 'Poppins', sans-serif;
        color: #4b3b76;
        border-radius: 5px;
        font-style: italic;
    }

    /* Star rating block */
    .star-rating {
        border-left: 4px solid #f5c518;  /* gold color border */
        background-color: #fff8dc;      /* light warm background */
        padding: 10px 15px;
        margin: 15px 0 10px 0;
        font-size: 26px;
        font-weight: 700;
        color: #f5c518;
        border-radius: 5px;
        font-family: 'Poppins', sans-serif;
        display: inline-block;
        user-select: none;
    }

    /* Export buttons container styling */
    .export-buttons > div {
        margin-top: 8px;
    }
</style>
""", unsafe_allow_html=True)

# --- Header ---
st.markdown("""
    <div class="header-footer typewriter">
        Welcome to Sentify Dashboard
    </div>
    <p style="font-family: 'Poppins', sans-serif; font-size: 32px; font-weight: 700; margin: 0; color: black; text-align: center;">
        Sentify Analysis
    </p>
    <p style="font-family: 'Poppins', sans-serif; font-size: 24px; font-weight: 700; margin: 0 0 20px 0; color: black; text-align: center;">
        Capturing emotional tones and translating them into star-rated insights
    </p>
""", unsafe_allow_html=True)

load_dotenv()
api_key = os.getenv("HUGGINGFACE_API_KEY")

# --- Theme Selector in Sidebar ---
theme_category = st.sidebar.selectbox(
    "🎨 Select Theme Category",
    options=["Plain Themes", "Nature Themes", "AI Animated Themes"]
)

if theme_category == "Plain Themes":
    theme = st.sidebar.selectbox(
        "Choose Plain Theme",
        options=["Minimal White", "Soft Blue", "Warm Yellow", "Cool Gray", "Elegant Black"]
    )
elif theme_category == "Nature Themes":
    theme = st.sidebar.selectbox(
        "Choose Nature Theme",
        options=["Mountain Landscape", "Ocean Waves", "Forest Canopy", "Sunset Sky", "Desert Dunes"]
    )
else:  # AI Animated Themes
    theme = st.sidebar.selectbox(
        "Choose AI Animated Theme",
        options=["AI Neural Network", "Digital Brain", "AI Processing", "Robot Vision"]
    )

# --- Dynamic Background Injection ---
theme_backgrounds = {
    # Plain Themes - Updated with proper color matching
    "Minimal White": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='100' height='100' viewBox='0 0 100 100'%3E%3Cg fill-opacity='0.03'%3E%3Cpath d='M50 0C22.4 0 0 22.4 0 50s22.4 50 50 50 50-22.4 50-50S77.6 0 50 0z'/%3E%3C/g%3E%3C/svg%3E",
    "Soft Blue": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='100' height='100' viewBox='0 0 100 100'%3E%3Cg fill='%236bb6ff' fill-opacity='0.1'%3E%3Cpath d='M50 0C22.4 0 0 22.4 0 50s22.4 50 50 50 50-22.4 50-50S77.6 0 50 0z'/%3E%3C/g%3E%3C/svg%3E",
    "Warm Yellow": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='100' height='100' viewBox='0 0 100 100'%3E%3Cg fill='%23FFE135' fill-opacity='0.15'%3E%3Cpath d='M50 0C22.4 0 0 22.4 0 50s22.4 50 50 50 50-22.4 50-50S77.6 0 50 0z'/%3E%3C/g%3E%3C/svg%3E",
    "Cool Gray": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='100' height='100' viewBox='0 0 100 100'%3E%3Cg fill='%23808080' fill-opacity='0.1'%3E%3Cpath d='M50 0C22.4 0 0 22.4 0 50s22.4 50 50 50 50-22.4 50-50S77.6 0 50 0z'/%3E%3C/g%3E%3C/svg%3E",
    "Elegant Black": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='100' height='100' viewBox='0 0 100 100'%3E%3Cg fill='%23000000' fill-opacity='0.05'%3E%3Cpath d='M50 0C22.4 0 0 22.4 0 50s22.4 50 50 50 50-22.4 50-50S77.6 0 50 0z'/%3E%3C/g%3E%3C/svg%3E",
    
    # Nature Themes
    "Mountain Landscape": "https://images.unsplash.com/photo-1449824913935-59a10b8d2000?ixlib=rb-4.0.3&auto=format&fit=crop&w=2070&q=80",
    "Ocean Waves": "https://images.unsplash.com/photo-1439066615861-d1af74d74000?ixlib=rb-4.0.3&auto=format&fit=crop&w=2070&q=80",
    "Forest Canopy": "https://images.unsplash.com/photo-1441974231531-c6227db76b6e?ixlib=rb-4.0.3&auto=format&fit=crop&w=2070&q=80",
    "Sunset Sky": "https://images.unsplash.com/photo-1506905925346-21bda4d32df4?ixlib=rb-4.0.3&auto=format&fit=crop&w=2070&q=80",
    "Desert Dunes": "https://images.unsplash.com/photo-1509316975850-ff9c5deb0cd9?ixlib=rb-4.0.3&auto=format&fit=crop&w=2070&q=80",
    
    # AI Animated GIF Themes - Updated with working calming tech GIFs
    "AI Neural Network": "https://i.gifer.com/7plX.gif",
    "Digital Brain": "https://i.gifer.com/VgGI.gif",
    "AI Processing": "https://i.gifer.com/UzE7.gif",
    "Robot Vision": "https://i.gifer.com/QBU.gif", 
}

# Set default theme to Desert Dunes
if 'theme_set' not in st.session_state:
    st.session_state.theme_set = "Desert Dunes"
    theme = "Desert Dunes"

background_image = theme_backgrounds[theme]

# Plain theme color styling
plain_theme_styles = {
    "Minimal White": {
        "background": "linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%)",
        "text_color": "#212529",
        "accent_color": "#6c757d"
    },
    "Soft Blue": {
        "background": "linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%)",
        "text_color": "#1565c0",
        "accent_color": "#2196f3"
    },
    "Warm Yellow": {
        "background": "linear-gradient(135deg, #fff9c4 0%, #fff59d 100%)",
        "text_color": "#f57f17",
        "accent_color": "#ffb300"
    },
    "Cool Gray": {
        "background": "linear-gradient(135deg, #f5f5f5 0%, #e0e0e0 100%)",
        "text_color": "#424242",
        "accent_color": "#757575"
    },
    "Elegant Black": {
        "background": "linear-gradient(135deg, #424242 0%, #212121 100%)",
        "text_color": "#ffffff",
        "accent_color": "#9e9e9e"
    }
}

# Apply theme-specific styling
if theme in plain_theme_styles:
    style = plain_theme_styles[theme]
    st.markdown(f"""
        <style>
        .stApp {{
            background: {style['background']} !important;
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
            min-height: 100vh;
        }}
        
        .header-footer {{
            background: {style['background']} !important;
            color: {style['text_color']} !important;
            border: 2px solid {style['accent_color']};
        }}
        
        .sentiment-box {{
            border-color: {style['accent_color']} !important;
            color: {style['text_color']} !important;
        }}
        
        .stButton > button {{
            background-color: {style['accent_color']} !important;
            color: {style['text_color']} !important;
        }}
        </style>
    """, unsafe_allow_html=True)
else:
    # For nature and AI themes, use the background image
    st.markdown(f"""
        <style>
        .stApp {{
            background-image: url('{background_image}');
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
            filter: brightness(0.85);
            min-height: 100vh;
        }}
        </style>
    """, unsafe_allow_html=True)

API_URL = "https://api-inference.huggingface.co/models/nlptown/bert-base-multilingual-uncased-sentiment"

def get_sentiment(text):
    headers = {"Authorization": f"Bearer {api_key}"}
    payload = {"inputs": text}
    response = requests.post(API_URL, headers=headers, json=payload)
    if response.status_code == 200:
        result = response.json()[0]
        sentiment_scores = {entry["label"]: entry["score"] for entry in result}
        # Sort scores to ensure consistent order for charts (e.g., 1-5 stars)
        sorted_scores = dict(sorted(sentiment_scores.items()))
        sentiment_label = max(sorted_scores, key=sorted_scores.get)
    else:
        # Improved error handling for API issues
        st.error(f"Error from Hugging Face API: {response.status_code} - {response.text}")
        sentiment_scores = {"Neutral": 0.5} # Fallback for error
        sentiment_label = "Neutral (API Error)"
    return sentiment_label, sentiment_scores

def generate_comment(sentiment_label):
    comments = {
        "1 star": "⚠️  Insight: Strong negativity observed. Consider revising your tone.",
        "2 stars": "😕  Insight: Some negativity detected. Consider neutral phrasing.",
        "3 stars": "😌  Insight: Neutral sentiment detected.",
        "4 stars": "😊  Insight: Positive sentiment observed!",
        "5 stars": "🌟  Insight: Highly positive sentiment!"
    }
    return comments.get(sentiment_label, "🤖 Insight: Sentiment analysis complete.")

def generate_star_rating(sentiment_label):
    stars_map = {
        "1 star": "⭐",
        "2 stars": "⭐⭐",
        "3 stars": "⭐⭐⭐",
        "4 stars": "⭐⭐⭐⭐",
        "5 stars": "⭐⭐⭐⭐⭐"
    }
    return stars_map.get(sentiment_label, "⭐")

def visualize_sentiment(sentiment_scores):
    # Ensure labels and values are ordered for consistent chart display
    ordered_labels = ["1 star", "2 stars", "3 stars", "4 stars", "5 stars"]
    values = [sentiment_scores.get(label, 0) for label in ordered_labels]
    
    color_map = {
        "5 stars": "#007BFF",   # Blue
        "4 stars": "#20c997",   # Teal
        "3 stars": "#ffc107",   # Yellow
        "2 stars": "#fd7e14",   # Orange
        "1 star":  "#dc3545"    # Red
    }
    # Ensure colors match the ordered labels
    colors = [color_map.get(label, "#888") for label in ordered_labels]

    # --- FIX for ValueError: remove hover_data when not using a DataFrame ---
    # Plotly Express automatically infers hover information from `names`, `values`, `x`, `y`.
    # We use hovertemplate for custom formatting which is the correct approach here.
    
    pie_chart = px.pie(
        names=ordered_labels,
        values=values,
        title="Sentiment Distribution",
        color=ordered_labels,
        color_discrete_map=color_map
    )
    # Update hover template to be more explicit if desired (optional, default is good)
    pie_chart.update_traces(hovertemplate="<b>%{label}</b><br>Score: %{value:.2f}<br>Percentage: %{percent:.1%}<extra></extra>")


    bar_chart = px.bar(
        x=ordered_labels,
        y=values,
        title="Sentiment Breakdown",
        labels={"x": "Rating", "y": "Score"},
        color=ordered_labels,
        color_discrete_map=color_map
    )
    bar_chart.update_traces(hovertemplate="<b>Rating: %{x}</b><br>Score: %{y:.2f}<extra></extra>")


    line_chart = px.line(
        x=ordered_labels,
        y=values,
        title="Sentiment Trend",
        markers=True,
        labels={"x": "Rating", "y": "Score"},
        color=px.Constant("Sentiment Score"), # Ensures consistent line color if not mapped to a varying column
        color_discrete_map={"Sentiment Score": "#2F3D79"} # Custom line color
    )
    line_chart.update_traces(hovertemplate="<b>Rating: %{x}</b><br>Score: %{y:.2f}<extra></extra>")

    return pie_chart, bar_chart, line_chart


def create_docx_bytes(sentiment_label, scores, comment, pie_chart, bar_chart, line_chart):
    doc = Document()
    doc.add_heading('Sentiment Analysis Report', 0)
    doc.add_paragraph(f"Sentiment: {sentiment_label}")
    doc.add_paragraph(f"AI Comment: {comment}")
    doc.add_heading("Scores:", level=1)
    for label, score in scores.items():
        doc.add_paragraph(f"{label}: {round(score, 3)}")

    # Save charts to bytes buffers
    pie_buf = io.BytesIO()
    bar_buf = io.BytesIO()
    line_buf = io.BytesIO()
    pie_chart.write_image(pie_buf, format='png')
    bar_chart.write_image(bar_buf, format='png')
    line_chart.write_image(line_buf, format='png')
    pie_buf.seek(0)
    bar_buf.seek(0)
    line_buf.seek(0)

    # Save images temporarily for docx (docx requires a file or BytesIO)
    pie_path = "pie_chart.png"
    bar_path = "bar_chart.png"
    line_path = "line_chart.png"
    with open(pie_path, "wb") as f:
        f.write(pie_buf.getbuffer())
    with open(bar_path, "wb") as f:
        f.write(bar_buf.getbuffer())
    with open(line_path, "wb") as f:
        f.write(line_buf.getbuffer())

    doc.add_picture(pie_path, width=Inches(5))
    doc.add_picture(bar_path, width=Inches(5))
    doc.add_picture(line_path, width=Inches(5))

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    # Cleanup temp files
    os.remove(pie_path)
    os.remove(bar_path)
    os.remove(line_path)

    return doc_bytes

def create_pdf_bytes(sentiment_label, scores, comment, pie_chart, bar_chart, line_chart):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 750, "Sentiment Analysis Report")
    c.setFont("Helvetica", 12)
    c.drawString(72, 730, f"Sentiment: {sentiment_label}")
    c.drawString(72, 710, f"AI Comment: {comment}")

    # List scores
    y = 690
    for label, score in scores.items():
        c.drawString(72, y, f"{label}: {round(score,3)}")
        y -= 20

    # Draw charts images
    pie_buf = io.BytesIO()
    bar_buf = io.BytesIO()
    line_buf = io.BytesIO()
    pie_chart.write_image(pie_buf, format='png')
    bar_chart.write_image(bar_buf, format='png')
    line_chart.write_image(line_buf, format='png')
    pie_buf.seek(0)
    bar_buf.seek(0)
    line_buf.seek(0)

    pie_img = ImageReader(pie_buf)
    bar_img = ImageReader(bar_buf)
    line_img = ImageReader(line_buf)

    c.drawImage(pie_img, 72, y-180, width=450, height=150)
    c.drawImage(bar_img, 72, y-350, width=450, height=150)
    c.drawImage(line_img, 72, y-520, width=450, height=150)

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

def create_json_string(sentiment_label, scores, comment):
    report = {
        "sentiment": sentiment_label,
        "scores": scores,
        "comment": comment
    }
    return json.dumps(report, indent=2)

def create_csv_string(sentiment_label, scores, comment):
    df = pd.DataFrame(list(scores.items()), columns=["Sentiment", "Score"])
    df.loc[len(df)] = ["Overall Sentiment", sentiment_label]
    df.loc[len(df)] = ["AI Comment", comment]
    return df.to_csv(index=False)

# --- App Layout ---
st.title("Sentify Dashboard")
st.write("Upload a file or enter text below to analyze its sentiment:")

# Initialize session state for text area if not present
if 'user_text_input' not in st.session_state:
    st.session_state.user_text_input = ""

uploaded_file = st.file_uploader("Upload a text file, PDF, or DOCX", type=["txt", "pdf", "docx"])

user_text = ""

if uploaded_file:
    # Set the user_text_input in session state for consistency
    if uploaded_file.type == "text/plain":
        user_text = uploaded_file.read().decode("utf-8")
    elif uploaded_file.type == "application/pdf":
        pdf_reader = PdfReader(uploaded_file)
        text_pages = [page.extract_text() for page in pdf_reader.pages]
        user_text = "\n".join(text_pages)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        paragraphs = [para.text for para in doc.paragraphs]
        user_text = "\n".join(paragraphs)
    else:
        st.warning("Unsupported file type.")
    st.session_state.user_text_input = user_text # Update session state with uploaded text
else:
    # Use session state for the text area to allow clearing
    st.session_state.user_text_input = st.text_area(
        "Or enter text to analyze manually:",
        value=st.session_state.user_text_input, # Bind value to session state
        height=150
    )
    user_text = st.session_state.user_text_input # Get current text from session state

col1, col2 = st.columns(2)

with col1:
    analyze_clicked = st.button("Analyze Sentiment")

with col2:
    # Clear All button
    if st.button("Clear All"):
        if "analysis" in st.session_state:
            del st.session_state.analysis # Clear analysis results
        if "export_dropdown" in st.session_state:
            st.session_state.export_dropdown = False # Collapse export dropdown
        st.session_state.user_text_input = "" # Clear text area
        st.rerun() # Use st.rerun() for clearing, as experimental_rerun is deprecated


if analyze_clicked and user_text.strip():
    # Check for API key before making the call
    if not api_key:
        st.error(
            """
            **Hugging Face API Key is missing!** 🔑
            Please set the `HUGGINGFACE_API_KEY` environment variable to use the sentiment analysis feature.

            **How to get and set your API Key:**
            1.  **Get a Token:**
                * Go to [Hugging Face](https://huggingface.co/settings/tokens).
                * Sign up or log in.
                * Create a new **User Access Token**. You can choose "read" role.
            2.  **Set as Environment Variable:**
                * **Local Development:** Create a file named `.env` in the same directory as your `app.py` and add:
                    ```
                    HUGGINGFACE_API_KEY="YOUR_HF_TOKEN_HERE"
                    ```
                    Replace `"YOUR_HF_TOKEN_HERE"` with the token you generated.
                * **Deployment (e.g., Streamlit Community Cloud):** Set it as a secret environment variable in your deployment settings.
            """
        )
    else:
        with st.spinner("Analyzing sentiment... This might take a moment for large texts."): # Add a spinner for user feedback
            label, scores = get_sentiment(user_text)
            comment = generate_comment(label)
            stars = generate_star_rating(label)
            pie, bar, line = visualize_sentiment(scores)

            # Store analysis in session state to reuse for export
            st.session_state.analysis = {
                "label": label,
                "scores": scores,
                "comment": comment,
                "stars": stars,
                "pie": pie,
                "bar": bar,
                "line": line,
            }

# Display results only if analysis is present in session state
if "analysis" in st.session_state and st.session_state.analysis:
    analysis = st.session_state.analysis
    st.markdown(f'<div class="sentiment-box">Overall Sentiment: {analysis["label"]}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="star-rating">{analysis["stars"]}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="ai-comment">{analysis["comment"]}</div>', unsafe_allow_html=True)
    st.plotly_chart(analysis["pie"], use_container_width=True)
    st.plotly_chart(analysis["bar"], use_container_width=True)
    st.plotly_chart(analysis["line"], use_container_width=True)

    # Generate exports
    pdf_bytes = create_pdf_bytes(analysis["label"], analysis["scores"], analysis["comment"],
                                 analysis["pie"], analysis["bar"], analysis["line"]).getvalue()
    docx_bytes = create_docx_bytes(analysis["label"], analysis["scores"], analysis["comment"],
                                   analysis["pie"], analysis["bar"], analysis["line"]).getvalue()
    json_str = create_json_string(analysis["label"], analysis["scores"], analysis["comment"])
    csv_str = create_csv_string(analysis["label"], analysis["scores"], analysis["comment"])

    if "export_dropdown" not in st.session_state:
        st.session_state.export_dropdown = False

    # Main dropdown toggle button (styled blue via CSS)
    if st.button("Export Report " + ("▼" if not st.session_state.export_dropdown else "▲")):
        st.session_state.export_dropdown = not st.session_state.export_dropdown

    # Export buttons shown only if dropdown toggled open
    if st.session_state.export_dropdown:
        with st.container():
            st.download_button(
                label="Download PDF",
                data=pdf_bytes,
                file_name="sentiment_report.pdf",
                mime="application/pdf",
                key="download_pdf",
                help="Download report as PDF"
            )
            st.download_button(
                label="Download DOCX",
                data=docx_bytes,
                file_name="sentiment_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx",
                help="Download report as DOCX"
            )
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name="sentiment_report.json",
                mime="application/json",
                key="download_json",
                help="Download report as JSON"
            )
            st.download_button(
                label="Download CSV",
                data=csv_str,
                file_name="sentiment_report.csv",
                mime="text/csv",
                key="download_csv",
                help="Download report as CSV"
            )

# --- "How it Works" / About Section (in sidebar) ---
st.sidebar.markdown("---") # Separator
with st.sidebar.expander("❓ How it Works"):
    st.markdown("""
    **Sentify Dashboard** leverages cutting-edge Artificial Intelligence to analyze the emotional tone of text.

    **Underlying Technology:**
    * **
    * **Hugging Face:** We use models hosted on the Hugging Face platform, a leading hub for open-source machine learning.
    * **BERT Model (`nlptown/bert-base-multilingual-uncased-sentiment`):** This specific model is a powerful pre-trained language model, fine-tuned for sentiment classification. It understands context and nuances in multiple languages.

    **What is Sentiment Analysis?**
    Sentiment analysis (also known as opinion mining) is the process of computationally identifying and categorizing opinions expressed in a piece of text, especially in order to determine whether the writer's attitude towards a particular topic, product, etc., is positive, negative, or neutral.

    **How Sentify Classifies:**
    The model analyzes your input text and assigns a probability score to each of five sentiment categories:
    * **1 star:** Very Negative
    * **2 stars:** Negative
    * **3 stars:** Neutral
    * **4 stars:** Positive
    * **5 stars:** Very Positive

    The category with the highest probability is then chosen as the overall sentiment.
    """)

# --- Footer ---
st.markdown("""
    <div class="header-footer footer">
        &copy; 2025 Sentify Dashboard. All rights reserved.
    </div>
""", unsafe_allow_html=True)
